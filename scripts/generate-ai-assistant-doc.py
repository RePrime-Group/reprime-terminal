#!/usr/bin/env python3
"""
Generate a concise stakeholder-facing DOCX describing the Terminal AI Assistant.

Run:
  python3 scripts/generate-ai-assistant-doc.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_PATH = "docs/ai-assistant/Terminal-AI-Assistant.docx"

PRIMARY_FONT = "Calibri"
HEADING_FONT = "Cambria"
COLOR_PRIMARY = RGBColor(0x0E, 0x34, 0x70)
COLOR_ACCENT = RGBColor(0xBC, 0x9C, 0x45)
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = PRIMARY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size in [(1, 18), (2, 13)]:
        h = styles[f"Heading {level}"]
        h.font.name = HEADING_FONT
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = COLOR_PRIMARY
        h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(4)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = HEADING_FONT
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = COLOR_PRIMARY
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.name = PRIMARY_FONT
    r2.font.size = Pt(12)
    r2.font.italic = True
    r2.font.color.rgb = COLOR_ACCENT


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = PRIMARY_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BODY


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            label, body = item
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            r1.font.name = PRIMARY_FONT
            r1.font.size = Pt(11)
            r1.font.color.rgb = COLOR_PRIMARY
            r2 = p.add_run(body)
            r2.font.name = PRIMARY_FONT
            r2.font.size = Pt(11)
            r2.font.color.rgb = COLOR_BODY
        else:
            r = p.add_run(item)
            r.font.name = PRIMARY_FONT
            r.font.size = Pt(11)
            r.font.color.rgb = COLOR_BODY


def build_document():
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    add_title(doc, "Terminal AI Assistant", "Product and Technical Overview")

    # 1. Overview
    doc.add_heading("Overview", level=1)
    add_paragraph(
        doc,
        "The Terminal AI Assistant is a chat experience inside the Reprime "
        "investor portal. It lets accredited investors ask questions about one "
        "specific deal and get answers grounded in real underwriting data. The "
        "assistant identifies itself as Terminal AI Assistant, speaks peer to "
        "peer, never invents numbers, and never discusses other deals."
    )

    # 2. Capabilities
    doc.add_heading("Capabilities", level=1)
    add_bullets(doc, [
        ("Deal facts", "cap rate, NOI, IRR, DSCR, cash on cash, equity, debt terms, exit cap, hold period."),
        ("Rent roll", "anchor tenant, top tenants by SF or income, lease expirations, WALT, concentration risk."),
        ("Documents", "list of due diligence files in the data room. Direct user to portal viewer for content."),
        ("Portfolios", "identifies the right building from the user's wording and scopes answers correctly."),
        ("What if questions", "directional reasoning anchored in base values, with a pointer to the Financial Modeling tab for exact numbers."),
        ("Format", "direct answer first, supporting numbers next, one relevant follow up question at the end."),
    ])

    # 3. How it works
    doc.add_heading("How It Works", level=1)
    add_paragraph(
        doc,
        "The user opens a deal, asks a question in the chat panel, and the "
        "assistant streams the answer back. Multi turn conversations are "
        "supported and the user can switch between past chat threads on the same "
        "deal."
    )
    add_paragraph(doc, "End to end flow per turn:")
    add_bullets(doc, [
        "Frontend posts to a Vercel API route which authenticates the session and forwards to n8n.",
        "n8n validates the request, confirms deal access, and looks up or creates the conversation record.",
        "The agent reads the system prompt and recent memory, then decides which tools to call.",
        "Each tool fetches only the columns needed via PostgREST, keeping token usage low.",
        "The agent composes the answer; the workflow saves the message and increments per user token usage.",
        "The final answer streams back to the frontend.",
    ])

    # 4. Tools
    doc.add_heading("Tools", level=1)
    add_paragraph(
        doc,
        "Four tools, all wired as HTTP request tools with column projection. "
        "The assistant selects only the columns it needs for each question."
    )
    add_bullets(doc, [
        ("Get Deal", "deal row with all financial metrics, status, dates, and document paths."),
        ("Get Addresses", "building list for portfolio deals; resolves user wording to the right building id."),
        ("Get Tenants", "rent roll with conditional building filter for portfolios."),
        ("Get Documents", "due diligence file list."),
    ])

    # 5. Architecture
    doc.add_heading("Architecture", level=1)
    add_bullets(doc, [
        ("Frontend", "Next.js on Vercel with React hooks for chat state and history."),
        ("API layer", "Vercel routes proxy to n8n, gating on Supabase Auth."),
        ("Orchestration", "n8n workflow named deal assistant chat."),
        ("Language model", "Claude Sonnet 4.6 via the Anthropic API."),
        ("Database", "Supabase PostgreSQL with row level security and PostgREST."),
        ("Memory", "n8n native chat memory, eight turn sliding window."),
    ])

    # 6. Security and observability
    doc.add_heading("Security and Observability", level=1)
    add_bullets(doc, [
        "Row level security on every chat table; investors see only their own data.",
        "Each conversation is locked to one deal; the assistant refuses cross deal questions and prompt injection attempts.",
        "Service role key is server side only, never bundled in the client.",
        "Real Anthropic token usage is captured per turn and accumulated per user via an atomic database function.",
        "Admin page at slash admin slash ai usage shows totals by user, by model, and recent activity.",
    ])

    # 7. Status
    doc.add_heading("Status", level=1)
    add_bullets(doc, [
        "All four tools migrated to HTTP request tools with column projection.",
        "System prompt finalized; identity is Terminal AI Assistant.",
        "Per user token accumulator live; admin page built.",
        "Workflow simplified; redundant nodes and columns removed.",
        "Ready for internal and beta usage.",
    ])

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Reprime Terminal, May 2026")
    r.font.name = HEADING_FONT
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COLOR_ACCENT

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT_PATH)
    print(f"✓ Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
